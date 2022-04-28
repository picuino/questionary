#
#  Program to convert questionnaries in YAML format to several other formats
#  like CSV spreadsheet, Moodle XML, json for JavaScript or Docx document.
#
#  Questionary Copyright (c) 2021 Carlos Pardo
#
#  This program is free software: you can redistribute it and/or modify
#  it under the terms of the GNU General Public License as published by
#  the Free Software Foundation, either version 3 of the License, or
#  (at your option) any later version.
#
#  This program is distributed in the hope that it will be useful,
#  but WITHOUT ANY WARRANTY; without even the implied warranty of
#  MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#  GNU General Public License for more details.
#
#  You should have received a copy of the GNU General Public License
#  along with this program.  If not, see <https://www.gnu.org/licenses/>.
#

import re
import os
import codecs
import base64
import random
import copy
import hashlib
import shutil

import yaml
import jinja2
from PIL import Image
import docx
from docx.shared import Inches, Cm, Pt


def main():
   """Main program"""

   moodle_template = 'moodle-multichoice-template.xml'
   json_template = 'json-template.json'
   html_template = 'game-template.html'

   multichoice_path = '.'
   build_path = 'build'
   html_path = '../docs'
   
   questionary = Questionary(overwrite=False)

   # Process all yaml files of this directory
   questions_counter = DictCounter()
   yaml_files = [yaml for yaml in os.listdir(multichoice_path) if yaml[-5:].lower() == '.yaml']
   for yaml_file in yaml_files:
      print('\nFile: %s/%s' % (multichoice_path, yaml_file))
      questions = questionary.read_yaml(yaml_file, path=multichoice_path)
      questions_counter.add(questions, yaml_file)
      questionary.write_csv(path=build_path)
      questionary.docx_generate(path=build_path)
      questionary.moodle_generate(moodle_template, path=build_path)
      questionary.json_generate(json_template, path=html_path)
      questionary.html_generate(html_template, path=html_path)
   print('\nTotal questions= %s' % str(questions_counter))



class DictCounter():
   def __init__(self):
     self.counter = {}

   def add(self, count, name):
      prefix = re.split('[_ -]', name, maxsplit=1)[0]
      if prefix in self.counter:
         self.counter[prefix] += count
      else:
         self.counter[prefix] = count
      
   def __str__(self):
      return str(self.counter)



class Questionary():

   def __init__(self, overwrite=False):
      self.templates_path = 'templates'
      self.px_cm = 59    # Pixel per centimeter. Used in docx images. 59 px_cm = 150dpi
      self.hash_len = 20
      self.csv_delimiter = ','
      self.questions = []
      self.overwrite = overwrite
      random.seed(1000)
      
      
   def read_yaml(self, filename, path='./'):
      """Read questions from Yaml file"""
      with codecs.open(os.path.join(path, filename), 'r', encoding='utf-8') as yamlfile:
         yamldata = yamlfile.read()
      yaml_files = list(yaml.load_all(yamldata, Loader=yaml.SafeLoader))
      self.header = yaml_files[0]
      self.questions = yaml_files[1]
      self.clean_questions()
      self.test_errors()
      self.yaml_path = path
      self.yaml_file = filename
      self.filename = os.path.splitext(os.path.basename(filename))[0]
      self.add_image_info()
      self.add_title()
      print('   Readed %d questions' % len(self.questions))
      return len(self.questions)


   def write_file(self, filename, data):
      if self.overwrite or self.file_older(filename):
         print('   Writing: ' + filename)
         if isinstance(data, docx.document.Document):
            data.save(filename)
         else:
            with codecs.open(filename, 'w', encoding='utf-8') as outfile:
               outfile.write(data)


   def file_older(self, filename1, filename2=False):
      if not filename2:
         filename2 = os.path.join(self.yaml_path, self.yaml_file)
      if os.path.exists(filename1):
         if os.path.getmtime(filename1) < os.path.getmtime(filename2):
            return True
         else:
            return False
      else:
         return True


   def not_key(self, dictionary, key):
      if key in dictionary and dictionary[key]:
         return False
      return True


   def clean_questions(self):
      for question in self.questions:
         new_text = question['Question'].strip('\n')
         if new_text != question['Question']:
            question['Question'] = new_text


   def test_errors(self):
      # Test errors in header
      if not 'Category' in self.header:
         print('   Warning: no Category value')
         self.header['Category'] = 'General'
      if not 'Copyright' in self.header:
         print('   Warning: no Copyright value')
         self.header['Copyright'] = ''
      if not 'Show_max' in self.header:
         print('   Warning: no Show_max value')
         self.header['Show_max'] = 0
      if not 'Title' in self.header:
         print('   Warning: no Title value')
         self.header['Title'] = 'No title available'

      # Test Questions
      safe_questions = []
      for question in self.questions:
         if not question:
            print('   Error: question empty ')
            continue
         if not 'Question' in question:
            print('   Error: question without text ' + str(question))
            continue
         if not 'Choices' in question:
            print('   Error: question without Choices ' + str(question))
            continue

         safe_questions.append(question)
      self.questions = safe_questions


   def add_title(self):
      """Add Title to every question if does not exists"""
      for question in self.questions:
         if not 'Title' in question or not question['Title']:
            question['Title'] = question['Question']
   
   
   def write_csv(self, path='./'):
      csv_filename = os.path.join(path, self.filename + '.csv')
      csv_data = ['Question;Image;Image_width;Choice_1;Choice_2;Choice_3;Choice_4;Choice_5;Choice_6;Block']
      for row in self.questions:
         line = '"' + row['Question'] + '";"'
         if row['Image']:
            line = line + str(row['Image']['filename']) + '";"'
            line = line + str(row['Image']['display_width']) + '";"'
         else:
            line = line + '";"";"'
         for i in range(6):
            if i < len(row['Choices']):
               line = line + str(row['Choices'][i]) + '";"'
            else:
               line = line + '";"'
         line = line + self.filename + '"'
         csv_data.append(line)
      csv_data = '\n'.join(csv_data)
      self.write_file(csv_filename, csv_data)
   
   
   def read_b64(self, filename):
      """Read image and returns data in ascii base64 format"""
      data = open(filename, 'rb').read()
      return base64.b64encode(data).decode('ascii')


   def hashname(self, filename):
      data = open(filename, 'rb').read()
      return hashlib.sha224(data).hexdigest()[:self.hash_len] + os.path.splitext(filename)[1]

   
   def add_image_info(self):
      """Read images from disk and add it several info and translations"""
      for question in self.questions:
         if 'Image' in question and question['Image']:
            imagedict = {}
            imagedict['filename'] = question['Image']
            if os.path.exists(imagedict['filename']):
               imagedict['hashname'] = self.hashname(imagedict['filename'])
               imagedict['path'] = os.path.dirname(imagedict['filename'])
               imagedict['mtime'] = os.path.getmtime(imagedict['filename'])
               imagedict['base64'] = self.read_b64(imagedict['filename'])
               width, height = Image.open(imagedict['filename']).size
               imagedict['width'] = width
               imagedict['height'] = height
               if not 'Image_width' in question:
                  imagedict['display_width'] = width
               else:
                  imagedict['display_width'] = question['Image_width']
                  del question['Image_width']
               if imagedict['display_width'] > 800:
                  print('   Warning image file too large. Display width=%3d  File=%s' % (imagedict['display_width'], imagedict['filename']))
               question['Image'] = imagedict                  
            else:
               print('Image file does not exists: ' + imagedict['filename'])
               question['Image'] = {}
         else:
            question['Image'] = {}
   
   
   def jinja_template(self, template_file):
      """Load jinja environment and template file.
         return a jinja template object"""
      templateLoader = jinja2.FileSystemLoader(searchpath=self.templates_path)
      templateEnv = jinja2.Environment(loader=templateLoader)
      self.template = templateEnv.get_template(template_file)
      
   
   def moodle_generate(self, template_file, path='./'):
      """Genera los cuestionarios en formato Moodle xml a partir de las
         cuestiones. Se genera un archivo por cada bloque de cuestiones"""
      xml_filename = os.path.join(path, self.filename + '.xml')
      self.jinja_template(template_file)
      xml_data = self.template.render(questions = self.questions, header = self.header, filename = self.filename)
      self.write_file(xml_filename, xml_data)
   
   
   def docx_make_head(self):
      self.docx = docx.Document()
   
      # Define page properties
      sections = self.docx.sections
      for section in sections:
          section.top_margin = Cm(1)
          section.bottom_margin = Cm(1)
          section.left_margin = Cm(1)
          section.right_margin = Cm(1)
      section = self.docx.sections[0]
      sectPr = section._sectPr
      cols = sectPr.xpath('./w:cols')[0]
      cols.set(docx.oxml.ns.qn('w:num'),'2')
   
      # Define Choice style
      style = self.docx.styles.add_style('Choice', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
      paragraph_format = style.paragraph_format
      paragraph_format.left_indent = Cm(1)
      paragraph_format.first_line_indent = Cm(-0.5)
      paragraph_format.line_spacing = 1
      paragraph_format.space_before = Pt(0)
      paragraph_format.space_after = Pt(0)
      tab_stops = paragraph_format.tab_stops
      tab_stops.add_tab_stop(Cm(4.75))
      tab_stops.add_tab_stop(Cm(5.25))
      paragraph_format.widow_control = True
   
      # Define Image style
      style = self.docx.styles.add_style('Image', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
      paragraph_format = style.paragraph_format
      paragraph_format.line_spacing = 1
      paragraph_format.space_before = Pt(2)
      paragraph_format.space_after = Pt(2)
      paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
   
      # Redefine List Number style
      style = self.docx.styles['List Number']
      paragraph_format = style.paragraph_format
      paragraph_format.space_before = Pt(12)
      paragraph_format.space_after = Pt(0)
      paragraph_format.left_indent = Cm(0.6)
      paragraph_format.first_line_indent = Cm(-0.6)
   
   
   def docx_add_image(self, question):
      """Add image to question in docx document"""
      image_width = question['Image']['display_width']
      par = self.docx.add_paragraph(style='Image')
      par.add_run().add_picture(question['Image']['filename'], width=Cm(image_width / self.px_cm))
   
   
   def suffle_choices(self, question):
       choices = copy.copy(question['Choices'])
       random.shuffle(choices)
       return choices
   
   
   def docx_add_questions(self):
      """Create docx questions"""
      self.docx.add_heading('%s. %s.' % (self.header['Category'], self.header['Title']), level=1)      
      for question in self.questions:      
         self.docx.add_paragraph(question['Question'], style='List Number')
         if question['Image']:
            self.docx_add_image(question)
         choices = self.suffle_choices(question)
         for i in range(len(choices)):
            self.docx.add_paragraph('abcdefg'[i] + ')\t' + str(choices[i]), style='Choice')
            
   
   def docx_generate(self, path='./'):
      """Genera un archivo docx con las preguntas y opciones de todas
         las cuestiones."""
      docx_filename = os.path.join(path, self.filename + '.docx')
      self.docx_make_head()
      self.docx_add_questions()
      self.write_file(docx_filename, self.docx)


   def json_generate(self, template_file, path='docs'):
      """Genera los archivos json a partir de las cuestiones."""
      # Copy json images
      for question in self.questions:
         if question['Image']:
            origin = question['Image']['filename']
            dest = os.path.join(path, 'images', question['Image']['hashname'])
            if self.file_older(dest, origin):
               print('   Writing: ' + origin)
               shutil.copy2(origin, dest)

      # Generate and save json
      json_filename = os.path.join(path, self.filename + '.json')
      self.jinja_template(template_file)
      json_data = self.template.render(questions = self.questions)
      self.write_file(json_filename, json_data)


   def html_generate(self, template_file, path='./'):
      """Genera los archivos html para jugar con las cuestiones."""
      html_filename = os.path.join(path, self.filename + '.html')
      self.jinja_template(template_file)
      html_data = self.template.render(filename=self.filename, header=self.header)
      self.write_file(html_filename, html_data)


if __name__ == "__main__":
   main()
